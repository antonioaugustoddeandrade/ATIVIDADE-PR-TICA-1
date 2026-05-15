from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def set_document_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12 if style_name != "Title" else 14)
        style.font.bold = True

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.revision = 1
    props.category = ""
    props.comments = ""
    props.identifier = ""
    props.keywords = ""
    props.language = "pt-BR"
    props.subject = ""
    props.version = ""


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_paragraph(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        rest = text[len(bold_label):]
        r2 = p.add_run(rest)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_manuscript():
    doc = Document()
    set_document_defaults(doc)
    add_page_number(doc.sections[0])
    doc.core_properties.title = "Esquizofrenia e psicoterapia integrada"

    # Page 1
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Esquizofrenia e psicoterapia integrada: revisão teórica recente e análise das práticas públicas do P.A.I. Papa Francisco, Hospital São Francisco na Providência de Deus (Tijuca, Rio de Janeiro)"
    )
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    add_paragraph(
        doc,
        "Tema: O artigo discute a esquizofrenia como transtorno psicótico de alta complexidade clínica e social, com ênfase em revisões teóricas e sistemáticas publicadas nos últimos cinco anos em português e espanhol, na literatura internacional recente sobre intervenções psicossociais e na descrição, a partir de fontes institucionais públicas, de recursos psicoterapêuticos e psicossociais oferecidos no P.A.I. Papa Francisco, serviço de saúde mental do Hospital São Francisco na Providência de Deus, na Tijuca, Rio de Janeiro.",
        "Tema:",
    )
    add_paragraph(
        doc,
        "Palavras-chave: esquizofrenia; psicoterapia; terapia cognitivo-comportamental; psicoeducação; reabilitação psicossocial; Hospital São Francisco na Providência de Deus.",
        "Palavras-chave:",
    )
    add_heading(doc, "Resumo", 1)
    add_paragraph(
        doc,
        "A esquizofrenia é uma condição psicótica grave, heterogênea e associada a sofrimento subjetivo, prejuízo funcional, estigma e necessidades complexas de cuidado. Este artigo apresenta uma revisão narrativa original sobre literatura recente em português, espanhol e inglês a respeito da esquizofrenia e de suas abordagens psicoterapêuticas, articulando esses achados com informações institucionais publicamente disponíveis sobre o P.A.I. Papa Francisco, serviço especializado em saúde mental do Hospital São Francisco na Providência de Deus, localizado na Tijuca, Rio de Janeiro. Foram priorizadas revisões teóricas, revisões sistemáticas, metanálises, diretrizes clínicas e páginas institucionais verificáveis. A literatura recente converge ao indicar que a psicoterapia não substitui o tratamento farmacológico quando este é clinicamente indicado, mas amplia o cuidado ao trabalhar sofrimento, adesão, prevenção de recaídas, funcionamento, cognição, família e reinserção social. Terapia cognitivo-comportamental para psicose, psicoeducação, intervenções familiares, treinamento de habilidades sociais, remediação cognitiva, terapias de terceira geração e estratégias metacognitivas aparecem como componentes relevantes, com níveis variáveis de evidência. No P.A.I., as informações públicas descrevem emergência psiquiátrica 24 horas, internação especializada, equipe multiprofissional, psicologia, terapia ocupacional, grupos terapêuticos, oficinas, acolhimento familiar e incentivo à continuidade pós-alta. Conclui-se que um modelo de psicoterapia para esquizofrenia em hospital geral deve ser integrado, ético, sensível ao estágio clínico e orientado à recuperação.",
    )
    page_break(doc)

    # Page 2
    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "Schizophrenia is a severe and heterogeneous psychotic condition associated with subjective distress, functional impairment, stigma, and complex care needs. This original narrative article reviews recent Portuguese, Spanish, and English literature on schizophrenia and psychotherapeutic approaches, and relates these findings to publicly available institutional information about P.A.I. Papa Francisco, the mental health service of Hospital São Francisco na Providência de Deus, located in Tijuca, Rio de Janeiro. The review prioritized theoretical reviews, systematic reviews, meta-analyses, clinical guidelines, and verifiable institutional pages. Recent evidence indicates that psychotherapy does not replace pharmacological treatment when clinically indicated, but broadens care by addressing distress, adherence, relapse prevention, functioning, cognition, family involvement, and social reintegration. Cognitive-behavioral therapy for psychosis, psychoeducation, family interventions, social skills training, cognitive remediation, third-wave therapies, and metacognitive strategies emerge as relevant components, with varying levels of evidence. Public information about P.A.I. describes 24-hour psychiatric emergency care, specialized hospitalization, a multidisciplinary team, psychology, occupational therapy, therapeutic groups, workshops, family support, and encouragement of post-discharge follow-up. The article concludes that psychotherapy for schizophrenia in a general hospital should be integrated, ethical, stage-sensitive, and recovery-oriented.",
    )
    add_heading(doc, "Resumen", 1)
    add_paragraph(
        doc,
        "La esquizofrenia es una condición psicótica grave y heterogénea, asociada con sufrimiento subjetivo, deterioro funcional, estigma y necesidades complejas de atención. Este artículo narrativo original revisa literatura reciente en portugués, español e inglés sobre esquizofrenia y abordajes psicoterapéuticos, relacionando esos hallazgos con información institucional pública del P.A.I. Papa Francisco, servicio de salud mental del Hospital São Francisco na Providência de Deus, ubicado en Tijuca, Río de Janeiro. Se priorizaron revisiones teóricas, revisiones sistemáticas, metaanálisis, guías clínicas y páginas institucionales verificables. La evidencia reciente indica que la psicoterapia no sustituye el tratamiento farmacológico cuando este está indicado, pero amplía el cuidado al abordar sufrimiento, adherencia, prevención de recaídas, funcionamiento, cognición, familia y reinserción social. La terapia cognitivo-conductual para psicosis, la psicoeducación, las intervenciones familiares, el entrenamiento en habilidades sociales, la rehabilitación cognitiva, las terapias de tercera generación y las estrategias metacognitivas aparecen como componentes relevantes, con niveles variables de evidencia. La información pública del P.A.I. describe urgencias psiquiátricas 24 horas, internación especializada, equipo multidisciplinario, psicología, terapia ocupacional, grupos terapéuticos, talleres, acogida familiar y continuidad posalta. Se concluye que la psicoterapia para esquizofrenia en hospital general debe ser integrada, ética, sensible a la etapa clínica y orientada a la recuperación.",
    )
    page_break(doc)

    # Page 3
    add_heading(doc, "Objetivo geral", 1)
    add_paragraph(
        doc,
        "Analisar, por meio de revisão narrativa de literatura recente e de fontes institucionais públicas, como o tratamento psicoterapêutico da esquizofrenia pode ser compreendido e organizado em articulação com práticas multiprofissionais descritas para o P.A.I. Papa Francisco, do Hospital São Francisco na Providência de Deus, na Tijuca, Rio de Janeiro.",
    )
    add_heading(doc, "Objetivos específicos", 1)
    add_bullet(
        doc,
        "Revisar contribuições teóricas recentes em português e espanhol sobre psicopatologia, cognição, sintomas positivos, alucinações, terapias cognitivas e terapias de terceira geração na esquizofrenia.",
    )
    add_bullet(
        doc,
        "Sintetizar evidências recentes de diretrizes, revisões sistemáticas e metanálises sobre intervenções psicoterapêuticas e psicossociais, com destaque para TCC para psicose, psicoeducação, intervenções familiares, remediação cognitiva e estratégias metacognitivas.",
    )
    add_bullet(
        doc,
        "Descrever, com base em informações públicas verificáveis, a estrutura, a equipe e as atividades terapêuticas divulgadas pelo P.A.I. Papa Francisco, considerando sua pertinência para pessoas com esquizofrenia ou outros transtornos psicóticos.",
    )
    add_bullet(
        doc,
        "Propor uma leitura ética e técnico-assistencial das práticas psicoterapêuticas possíveis no contexto hospitalar da Tijuca, distinguindo dados publicados de inferências clínicas que dependeriam de confirmação institucional, entrevista ou pesquisa aprovada por comitê de ética.",
    )
    add_heading(doc, "Delimitação metodológica", 1)
    add_paragraph(
        doc,
        "Trata-se de artigo de revisão narrativa e análise documental de fontes abertas. Não foram acessados prontuários, entrevistas, dados pessoais, imagens de pacientes ou informações internas do hospital. Assim, quando o texto aborda o P.A.I., refere-se exclusivamente ao que a instituição divulga em páginas públicas e ao que pode ser inferido, com cautela, à luz da literatura científica. Essa delimitação preserva confidencialidade e evita transformar dados institucionais gerais em afirmações diagnósticas sobre pessoas atendidas.",
    )
    page_break(doc)

    # Page 4
    add_heading(doc, "Introdução", 1)
    add_paragraph(
        doc,
        "A esquizofrenia costuma ser apresentada como um transtorno mental grave, marcado por experiências psicóticas, alterações cognitivas, sintomas negativos e impactos no funcionamento social. A Organização Mundial da Saúde estima que aproximadamente 24 milhões de pessoas vivam com esquizofrenia no mundo, chamando atenção para incapacidade, morte prematura, violações de direitos e lacunas de acesso a cuidado especializado (World Health Organization, 2022). A relevância clínica do tema, portanto, não se limita à descrição de delírios ou alucinações. Ela envolve trajetória de vida, estigma, vínculos, autonomia, adesão terapêutica, condições socioeconômicas e a capacidade dos serviços de saúde de oferecer respostas contínuas e humanizadas.",
    )
    add_paragraph(
        doc,
        "Nas últimas décadas, a compreensão do transtorno se deslocou de explicações exclusivamente biologizantes para modelos que integram neurodesenvolvimento, vulnerabilidade, cognição, trauma, fatores sociais e experiência subjetiva. Em língua espanhola, Masedo Gutiérrez (2021) revisou aspectos epistemológicos e fenomenológico-comportamentais da psicopatologia da esquizofrenia, criticando reducionismos e defendendo uma clínica capaz de escutar a alteração do self, da corporalidade e do estar-no-mundo. Pena-Garijo e Monfort-Escrig (2020), por sua vez, enfatizaram cognição social, metacognição e vieses de raciocínio, mostrando que a reabilitação não pode se restringir à remissão sintomática.",
    )
    add_paragraph(
        doc,
        "Em português, revisões recentes também convergem para a ampliação do cuidado. Gomes e Fernandes (2022) examinaram terapias de terceira geração nas perturbações psicóticas, destacando mindfulness, terapia de aceitação e compromisso e terapia focada na compaixão. Freitas e Valadas (2020) revisaram a abordagem cognitivo-comportamental das alucinações auditivo-verbais, descrevendo seu foco na relação da pessoa com as vozes e nas crenças associadas a elas. Vilar, Nogueira, Valentim e Seabra (2020) mapearam a psicoeducação como intervenção associada à adesão terapêutica em pessoas com esquizofrenia. Essas contribuições dialogam com diretrizes e metanálises internacionais, que sustentam a necessidade de planos abrangentes e centrados na pessoa (American Psychiatric Association, 2020; McDonagh et al., 2022).",
    )
    page_break(doc)

    # Page 5
    add_heading(doc, "Introdução (continuação)", 1)
    add_paragraph(
        doc,
        "O recorte institucional deste artigo é o P.A.I. Papa Francisco, descrito publicamente como Polo de Atenção Integral à Saúde Mental do Hospital São Francisco na Providência de Deus. A instituição informa que o serviço foi inaugurado em 2013, localiza-se na principal via da Tijuca e oferece atendimento a necessidades psiquiátricas, dependência química e outros transtornos, com emergência psiquiátrica 24 horas, internação, equipe multiprofissional e recursos terapêuticos como grupos, oficinas, acolhimento familiar e atividades físicas (Hospital São Francisco na Providência de Deus, n.d.-a, n.d.-b). Embora as páginas públicas não apresentem um protocolo específico para esquizofrenia, seus componentes são compatíveis com recomendações contemporâneas para transtornos psicóticos graves: cuidado multiprofissional, continuidade pós-alta, inclusão da família, redução de isolamento e intervenções psicossociais estruturadas.",
    )
    add_paragraph(
        doc,
        "A pergunta que orienta o desenvolvimento é: de que modo a literatura teórica recente e as evidências sobre psicoterapia em esquizofrenia ajudam a compreender as práticas públicas de cuidado em saúde mental descritas para o P.A.I. Papa Francisco? A resposta exige prudência. Não se pode afirmar, sem pesquisa de campo, quais técnicas são efetivamente aplicadas a cada paciente, sua frequência, fidelidade metodológica ou resultados. Ainda assim, é possível examinar como grupos terapêuticos, oficinas de arte, atividades de autocuidado, acolhimento da família, psicologia e terapia ocupacional podem compor um plano psicoterapêutico ampliado, desde que integrados a avaliação clínica, manejo medicamentoso, proteção de direitos e projetos de reabilitação.",
    )
    add_paragraph(
        doc,
        "Este texto, portanto, não é um relato de caso nem estudo avaliativo do hospital. É uma revisão narrativa original com finalidade acadêmica, construída a partir de literatura recente e documentos públicos. Ao fazê-lo, preserva o sigilo assistencial e oferece uma leitura crítica do tratamento psicoterapêutico da esquizofrenia em contexto hospitalar geral, destacando tanto possibilidades quanto limites.",
    )
    page_break(doc)

    # Page 6
    add_heading(doc, "Desenvolvimento: esquizofrenia, subjetividade e psicopatologia", 1)
    add_paragraph(
        doc,
        "A tradição psicopatológica contemporânea tem questionado leituras que reduzem a esquizofrenia a uma lista de sintomas positivos e negativos. A revisão de Masedo Gutiérrez (2021) recupera aportes fenomenológicos para compreender a esquizofrenia como alteração da ipseidade, da corporalidade e da evidência natural do mundo. Nessa leitura, delírios e alucinações não são eventos isolados, mas experiências que emergem em uma transformação mais ampla do modo como a pessoa se sente autora de pensamentos, ações e percepções. O sofrimento psicótico envolve estranhamento de si, ruptura de continuidade, perda de familiaridade com o cotidiano e dificuldade de partilhar significados.",
    )
    add_paragraph(
        doc,
        "Esse ponto é clinicamente importante porque uma psicoterapia orientada apenas a contestar crenças delirantes pode aumentar defensividade ou humilhação. Abordagens atuais recomendam formulação compartilhada, validação do sofrimento, investigação de gatilhos e construção de estratégias de manejo. A pessoa não é tratada como portadora de uma crença absurda a ser corrigida, mas como alguém que tenta compreender experiências ameaçadoras. Essa postura é coerente com a TCC para psicose, com terapias contextuais e com intervenções focadas em recuperação.",
    )
    add_paragraph(
        doc,
        "Outro eixo teórico é a cognição. Pena-Garijo e Monfort-Escrig (2020) descrevem déficits neurocognitivos, cognição social, metacognição e vieses como salto a conclusões, atribuições externalizantes e resistência a evidências desconfirmatórias. Esses fatores influenciam sintomas, funcionamento e relações sociais. Programas como treinamento metacognitivo e reabilitação cognitiva buscam reduzir vieses, melhorar reflexão sobre pensamentos e favorecer participação social. Essa perspectiva desloca o foco de uma psicoterapia meramente verbal para intervenções estruturadas, graduais e frequentemente grupais, nas quais o paciente pratica habilidades que se generalizam para a vida diária.",
    )
    page_break(doc)

    # Page 7
    add_heading(doc, "Revisões recentes em português e espanhol", 1)
    add_paragraph(
        doc,
        "Entre as revisões em português, Vilar et al. (2020) identificaram a psicoeducação como recurso para adesão terapêutica em pessoas com esquizofrenia. O artigo destacou intervenções com informações sobre transtorno, medicação e estratégias de coping, algumas conduzidas por enfermeiros de saúde mental. O resultado central foi aumento de conhecimento e mudança de atitude diante do regime medicamentoso. Embora a revisão tenha incluído poucos estudos, seu valor está em reforçar que adesão não depende apenas de prescrição, mas de compreensão, relação terapêutica, manejo de efeitos adversos, insight e participação ativa do paciente.",
    )
    add_paragraph(
        doc,
        "Freitas e Valadas (2020) concentraram-se nas alucinações auditivo-verbais. A TCC para vozes procura modificar crenças sobre poder, intenção e controle das vozes, reduzindo medo, submissão e comportamentos de evitação. Em vez de prometer eliminação completa das alucinações, busca reduzir sofrimento e incapacidade. Essa distinção é essencial em contextos hospitalares, pois pacientes em crise podem necessitar de estabilização medicamentosa e ambiental antes de intervenções cognitivas mais elaboradas. Em fases de maior estabilidade, a psicoterapia pode trabalhar diário de vozes, identificação de padrões, respostas alternativas e reconstrução de sentido.",
    )
    add_paragraph(
        doc,
        "Gomes e Fernandes (2022) revisaram terapias de terceira geração nas perturbações psicóticas. Mindfulness, aceitação e compromisso e compaixão não visam convencer a pessoa de que a experiência psicótica é falsa, mas alterar sua relação com eventos internos, reduzir fusão cognitiva, cultivar autocompaixão e ampliar ações orientadas por valores. Em espanhol, Torres Hernández, González Lorenzo e Martín Estévez (2021) revisaram técnicas cognitivas e reconheceram resultados favoráveis, ainda que não conclusivos. Martín López (2024) revisou intervenções de terceira geração e outras psicoterapias para sintomas positivos, incluindo terapia metacognitiva, aceitação e compromisso e AVATAR, apontando achados promissores com necessidade de amostras maiores e maior rigor.",
    )
    page_break(doc)

    # Page 8
    add_heading(doc, "Evidência internacional e diretrizes clínicas", 1)
    add_paragraph(
        doc,
        "A diretriz da American Psychiatric Association (2020) recomenda plano abrangente, documentado, centrado na pessoa e composto por tratamentos farmacológicos e não farmacológicos baseados em evidências. Entre as intervenções psicossociais, inclui TCC para psicose, psicoeducação, emprego apoiado e cuidado coordenado para primeiro episódio psicótico. A ênfase não está em uma técnica isolada, mas na combinação coerente entre avaliação, preferências do paciente, metas funcionais e continuidade.",
    )
    add_paragraph(
        doc,
        "McDonagh et al. (2022) atualizaram revisões sobre intervenções psicossociais para adultos com esquizofrenia e identificaram benefícios em funcionamento, qualidade de vida, sintomas e recaídas. Intervenções familiares, psicoeducação, manejo da doença, TCC, treinamento de habilidades sociais, emprego apoiado e serviços para primeiro episódio psicótico apresentaram efeitos em diferentes desfechos, ainda que a força da evidência tenha sido frequentemente baixa a moderada. Solmi et al. (2023), em revisão guarda-chuva de metanálises, também indicaram que serviços devem priorizar intervenções familiares e de início precoce em psicose inicial e, na esquizofrenia estabelecida, TCC, remediação cognitiva, intervenções familiares e emprego apoiado, observando heterogeneidade e qualidade variável da evidência.",
    )
    add_paragraph(
        doc,
        "A TCC, apesar de central, deve ser apresentada de modo equilibrado. Berendsen, Berendse, van der Torren, Vermeulen e de Haan (2024) encontraram evidência sugestiva, mas não convincente, para efeitos da TCC em psicopatologia geral, delírios e alucinações no fim do tratamento, com evidência fraca ou não significativa para sintomas negativos e sustentabilidade limitada no seguimento. Isso não invalida a TCC; apenas exige indicação realista, integração com outros cuidados e avaliação de objetivos concretos, como sofrimento, coping, funcionamento e prevenção de recaídas.",
    )
    page_break(doc)

    # Page 9
    add_heading(doc, "Família, psicoeducação e prevenção de recaídas", 1)
    add_paragraph(
        doc,
        "A esquizofrenia repercute sobre famílias e cuidadores, que muitas vezes enfrentam medo, sobrecarga, desinformação e conflitos comunicacionais. Rodolico et al. (2022), em revisão sistemática e metanálise em rede, compararam modelos de intervenção familiar para prevenção de recaídas e concluíram que quase todos reduziram recaídas em comparação ao tratamento usual. A psicoeducação familiar isolada mostrou desempenho favorável, inclusive em contextos com restrições de recursos. Esse achado é particularmente relevante para hospitais gerais, pois a internação pode ser uma janela para orientar família, reduzir culpa, planejar alta, mapear sinais precoces e estabelecer rotas de cuidado.",
    )
    add_paragraph(
        doc,
        "A psicoeducação não deve ser confundida com palestra unidirecional. Ela inclui escuta, adaptação cultural, linguagem simples, negociação de metas e atenção às experiências do paciente. Conteúdos mínimos abrangem natureza do transtorno, sinais de alerta, medicação e efeitos adversos, sono, substâncias psicoativas, manejo de estresse, prevenção de recaídas, direitos, rede de atenção psicossocial e crise. Em pessoas com esquizofrenia, a informação precisa ser dosada conforme estado clínico, cognição, nível de desorganização e disponibilidade emocional.",
    )
    add_paragraph(
        doc,
        "No contexto do P.A.I., a página institucional informa a existência de acolhimento da família, definido como conversas em que familiares podem trocar experiências e minimizar angústia e ansiedade pela recuperação dos pacientes (Hospital São Francisco na Providência de Deus, n.d.-a). Essa prática é compatível com recomendações de intervenção familiar, mas o documento público não informa se há protocolo manualizado, número de sessões ou avaliação de resultados. A leitura prudente é reconhecer o componente familiar como recurso terapêutico potencial, cuja efetividade dependerá de estrutura, frequência, formação da equipe e integração ao plano individual.",
    )
    page_break(doc)

    # Page 10
    add_heading(doc, "Psicoterapia no hospital geral: fases e objetivos", 1)
    add_paragraph(
        doc,
        "A psicoterapia da esquizofrenia em hospital geral precisa respeitar fases clínicas. Na crise aguda, objetivos prioritários são segurança, vínculo, redução de estímulos desorganizadores, manejo de risco, orientação breve e coordenação com psiquiatria e enfermagem. Intervenções longas e interpretativas podem ser inadequadas quando há intensa desorganização, agitação, risco suicida ou intoxicação. Nessa fase, a psicoterapia pode ocorrer como acolhimento estruturado, validação do medo, técnicas de grounding, respiração, identificação de necessidades imediatas e contato com família.",
    )
    add_paragraph(
        doc,
        "Durante estabilização, tornam-se possíveis formulação compartilhada, psicoeducação, análise de eventos precipitantes, planejamento de alta e prevenção de recaídas. O terapeuta pode trabalhar sinais precoces, rotina de sono, uso de substâncias, estressores familiares, adesão a consultas, estratégias para lidar com vozes e redução de isolamento. Em etapa de reabilitação, ganham importância habilidades sociais, cognição, autocuidado, projetos de estudo ou trabalho, atividades significativas e pertencimento comunitário.",
    )
    add_paragraph(
        doc,
        "O serviço descrito pelo P.A.I. possui características que favorecem essa lógica de fases: emergência psiquiátrica 24 horas, leitos de estabilização e urgência, alas de internação, equipe com médicos, psicólogas, terapeutas ocupacionais, assistente social, educador físico, nutricionista e enfermagem, além de espaços para oficinas e grupos (Hospital São Francisco na Providência de Deus, n.d.-a, n.d.-b). A presença em hospital geral também é relevante para manejo de comorbidades clínicas. Pessoas com esquizofrenia apresentam maior risco de condições metabólicas, cardiovasculares e infecciosas; portanto, a integração clínica reduz fragmentação e amplia segurança.",
    )
    page_break(doc)

    # Page 11
    add_heading(doc, "Práticas públicas do P.A.I. Papa Francisco e pertinência psicoterapêutica", 1)
    add_paragraph(
        doc,
        "A página do P.A.I. informa que o serviço oferece atendimento médico e psicológico, internação com suporte nutricional, farmacêutico e atenção especializada, além de incentivo à continuidade após alta hospitalar (Hospital São Francisco na Providência de Deus, n.d.-a). Também descreve salas para oficinas terapêuticas, sala de artes, sala multiuso, biblioteca, sala de beleza e atividades físicas. Esses recursos não são meros complementos recreativos quando articulados a um projeto terapêutico: podem funcionar como dispositivos de expressão, treino de rotina, estimulação cognitiva, recuperação de autonomia, ampliação de repertório social e reconstrução de autoestima.",
    )
    add_paragraph(
        doc,
        "Oficinas de artes e trabalhos manuais podem favorecer atenção, planejamento, coordenação motora, expressão simbólica e experiência de competência. A sala de beleza, descrita como promotora de autonomia, autocuidado e autoestima, pode ser relevante para sintomas negativos, empobrecimento de iniciativa e retraimento social. Atividades físicas com educador físico, três vezes por semana segundo a instituição, podem contribuir para bem-estar, sono, saúde metabólica e engajamento. Oficinas de nutrição e preparo de receitas saudáveis dialogam com o cuidado de efeitos metabólicos de antipsicóticos e com a educação em saúde.",
    )
    add_paragraph(
        doc,
        "O grupo de saúde mental, conduzido pela Psicologia e Terapia Ocupacional, é descrito como espaço semanal para promover interação social, inclusão, reflexão, música, autoestima, autonomia, relaxamento e respiração (Hospital São Francisco na Providência de Deus, n.d.-a). Para pessoas com esquizofrenia, tal grupo pode operar como intervenção psicoterapêutica de baixa intensidade, orientada a socialização, regulação emocional e prática de habilidades. Contudo, para caracterizá-lo como TCC, ACT, treinamento metacognitivo ou outra modalidade específica, seria necessário acesso a manual, objetivos, formação dos facilitadores e indicadores de processo.",
    )
    page_break(doc)

    # Page 12
    add_heading(doc, "Articulação entre evidência e prática no P.A.I.", 1)
    add_paragraph(
        doc,
        "A literatura sugere que um programa psicoterapêutico robusto para esquizofrenia no P.A.I. poderia organizar-se em quatro eixos integrados. O primeiro é estabilização e vínculo: acolhimento, avaliação de risco, redução de medo e construção de aliança. O segundo é psicoeducação individual e familiar: explicação do transtorno, sinais de recaída, medicação, efeitos adversos, sono, drogas e continuidade de cuidado. O terceiro é intervenção cognitivo-comportamental e metacognitiva: manejo de vozes, questionamento colaborativo de interpretações ameaçadoras, redução de vieses e aumento de flexibilidade. O quarto é reabilitação psicossocial: oficinas, autocuidado, atividade física, habilidades sociais, projetos de vida e rede comunitária.",
    )
    add_paragraph(
        doc,
        "Esses eixos podem ser aplicados em intensidade variável. Nem toda pessoa internada com esquizofrenia se beneficia de terapia cognitiva formal durante crise; algumas necessitam inicialmente de contenção ambiental, medicação, sono e orientação breve. Outras, já estabilizadas, podem avançar em formulação de caso, estratégias para alucinações e planejamento de reinserção. A personalização é um princípio ético e técnico: psicoterapia não é pacote uniforme, mas resposta graduada às necessidades, preferências e riscos.",
    )
    add_paragraph(
        doc,
        "Parra-Bolaños, Martínez Suárez e Velásquez Pérez (2022) defendem abordagens contemporâneas e transdisciplinares para diagnóstico e intervenção psiquiátrica na esquizofrenia. Esse argumento se aproxima do modelo institucional do P.A.I., que divulga equipe multiprofissional e cuidado integral. Ainda assim, a transdisciplinaridade só se realiza quando há plano compartilhado, comunicação entre profissionais, registros claros, participação do paciente e revisão de resultados. Sem esses elementos, múltiplas atividades podem coexistir sem produzir continuidade terapêutica.",
    )
    page_break(doc)

    # Page 13
    add_heading(doc, "Questões éticas, autoria e submissão", 1)
    add_paragraph(
        doc,
        "Este manuscrito foi elaborado como contribuição original de revisão narrativa e análise de fontes públicas. Não foi submetido simultaneamente a outra revista. Caso os autores venham a submeter versão derivada a periódico científico, recomenda-se declarar no sistema de submissão que o texto é original, inédito, não está em avaliação em outro periódico e que eventuais conflitos de interesse, financiamentos e contribuições individuais foram informados com transparência.",
    )
    add_paragraph(
        doc,
        "Quanto aos procedimentos éticos, este artigo não envolve participantes humanos identificáveis, prontuários, entrevistas, imagens ou coleta de dados sensíveis. As informações institucionais sobre o hospital foram extraídas de páginas públicas. Se, em etapa futura, os autores desejarem transformar o tema em pesquisa empírica no P.A.I., com profissionais, pacientes, familiares ou documentos internos, será necessário submeter protocolo a Comitê de Ética em Pesquisa, obter anuências institucionais, termos de consentimento quando aplicáveis e plano de proteção de dados.",
    )
    add_paragraph(
        doc,
        "Para avaliação por pares em modalidade duplo-cega, o arquivo principal deve permanecer sem nomes de autores, afiliações, agradecimentos identificáveis e metadados de autoria. A folha de rosto e o checklist de submissão devem ser enviados como documentos suplementares, conforme diretrizes da revista. Após a submissão, a sequência de autores informada na folha de rosto tende a ser adotada na publicação e alterações de autoria podem não ser permitidas, salvo regras específicas do periódico. Como a revista não foi indicada neste pedido, os modelos suplementares anexos a este pacote devem ser ajustados ao formulário oficial da revista escolhida.",
    )
    page_break(doc)

    # Page 14
    add_heading(doc, "Conclusão", 1)
    add_paragraph(
        doc,
        "A esquizofrenia exige cuidado que reconheça a realidade do sofrimento psicótico sem reduzir a pessoa ao diagnóstico. Revisões recentes em português e espanhol mostram uma clínica cada vez mais atenta à subjetividade, às vozes, à cognição, à metacognição, às terapias de terceira geração, à psicoeducação e ao funcionamento social. As evidências internacionais reforçam que intervenções psicossociais têm papel complementar essencial, especialmente TCC para psicose, psicoeducação, intervenções familiares, remediação cognitiva, treinamento de habilidades sociais, emprego apoiado e serviços de primeiro episódio. Os efeitos são heterogêneos e dependem de qualidade metodológica, contexto, fidelidade e metas realistas.",
    )
    add_paragraph(
        doc,
        "No Hospital São Francisco na Providência de Deus, o P.A.I. Papa Francisco divulga estrutura compatível com cuidado integral: emergência psiquiátrica, internação especializada, equipe multiprofissional, psicologia, terapia ocupacional, oficinas, grupos, atividades físicas e acolhimento familiar. Para pessoas com esquizofrenia, esses recursos podem compor tratamento psicoterapêutico ampliado quando articulados em plano individual, respeitando fases clínicas, direitos, autonomia e continuidade pós-alta. A prudência científica exige reconhecer que as fontes públicas não demonstram protocolos específicos nem resultados clínicos; elas apenas permitem identificar dispositivos potencialmente psicoterapêuticos e reabilitadores.",
    )
    add_paragraph(
        doc,
        "O desafio contemporâneo é transformar a internação, muitas vezes vivida como ruptura, em oportunidade de reconstrução de vínculo, compreensão, segurança e projeto de vida. Uma psicoterapia ética da esquizofrenia no hospital geral deve ser colaborativa, antistigmatizante, culturalmente sensível e integrada ao território. Quando o cuidado combina estabilização clínica, escuta, família, grupos, oficinas, manejo de sintomas e reabilitação, a meta deixa de ser apenas controlar a crise e passa a incluir recuperação possível, pertencimento e dignidade.",
    )
    page_break(doc)

    # Page 15
    add_heading(doc, "Referências", 1)
    references = [
        "American Psychiatric Association. (2020). The American Psychiatric Association practice guideline for the treatment of patients with schizophrenia (3rd ed.). American Psychiatric Association Publishing. https://doi.org/10.1176/appi.books.9780890424841",
        "Berendsen, S., Berendse, S., van der Torren, J., Vermeulen, J., & de Haan, L. (2024). Cognitive behavioural therapy for the treatment of schizophrenia spectrum disorders: An umbrella review of meta-analyses of randomised controlled trials. eClinicalMedicine, 67, Article 102392. https://doi.org/10.1016/j.eclinm.2023.102392",
        "Bighelli, I., Wallis, S., Reitmeir, C., Schwermann, F., Salahuddin, N. H., & Leucht, S. (2023). Effects of psychological treatments on functioning in people with schizophrenia: A systematic review and meta-analysis of randomized controlled trials. European Archives of Psychiatry and Clinical Neuroscience, 273, 779-810. https://doi.org/10.1007/s00406-022-01526-1",
        "Freitas, R. M., & Valadas, M. T. (2020). Alucinações auditivo-verbais na esquizofrenia: Uma revisão sobre a abordagem cognitivo-comportamental. Psilogos, 18(1-2). https://revistas.rcaap.pt/psilogos/article/view/19469",
        "Gomes, M., & Fernandes, N. (2022). Terapias de terceira geração nas perturbações psicóticas. Revista Portuguesa de Psiquiatria e Saúde Mental, 8(2), 58-65. https://doi.org/10.51338/rppsm.319",
        "Hospital São Francisco na Providência de Deus. (n.d.-a). P.A.I. Papa Francisco: Polo de Atenção Integral à Saúde Mental. Retrieved May 15, 2026, from https://hospitalsaofranciscorj.com.br/pai/",
        "Hospital São Francisco na Providência de Deus. (n.d.-b). Emergência e internação psiquiátricas - P.A.I. Retrieved May 15, 2026, from https://hospitalsaofranciscorj.com.br/emergencia-e-internacao-psiquiatricas/",
        "Martín López, A. (2024). Intervenciones cognitivo-conductuales de tercera generación y otras psicoterapias para tratar los síntomas positivos en la esquizofrenia: Una revisión teórica sistemática [Master's thesis, Universidad de Alcalá]. e_Buah. http://hdl.handle.net/10017/61608",
        "Masedo Gutiérrez, A. I. (2021). La psicopatología de la esquizofrenia. Revista de la Asociación Española de Neuropsiquiatría, 41(140). https://dx.doi.org/10.4321/s0211-57352021000200004",
        "McDonagh, M. S., Dana, T., Kopelovich, S. L., Monroe-DeVita, M., Blazina, I., Bougatsos, C., Grusing, S., & Selph, S. S. (2022). Psychosocial interventions for adults with schizophrenia: An overview and update of systematic reviews. Psychiatric Services, 73(3), 299-312. https://doi.org/10.1176/appi.ps.202000649",
        "Parra-Bolaños, N., Martínez Suárez, P. C., & Velásquez Pérez, L. A. (2022). Abordajes y tratamientos contemporáneos para la esquizofrenia: Diagnóstico e intervención psiquiátrica. Ciencia Latina Revista Científica Multidisciplinar, 6(6), 13221-13238. https://doi.org/10.37811/cl_rcm.v6i6.4324",
        "Pena-Garijo, J., & Monfort-Escrig, C. (2020). Cognición en la esquizofrenia. Estado actual de la cuestión (II): Sesgos cognitivos, modelos explicativos y programas de intervención. Revista de la Asociación Española de Neuropsiquiatría, 40(137). https://dx.doi.org/10.4321/s0211-57352020000100008",
        "Rodolico, A., Bighelli, I., Avanzato, C., Concerto, C., Cutrufelli, P., Mineo, L., Schneider-Thoma, J., Siafis, S., Signorelli, M. S., Wu, H., Wang, D., Furukawa, T. A., Pitschel-Walz, G., Aguglia, E., & Leucht, S. (2022). Family interventions for relapse prevention in schizophrenia: A systematic review and network meta-analysis. The Lancet Psychiatry, 9(3), 211-221. https://doi.org/10.1016/S2215-0366(21)00437-5",
        "Solmi, M., Croatto, G., Piva, G., Rosson, S., Fusar-Poli, P., Rubio, J. M., Carvalho, A. F., Vieta, E., Arango, C., DeTore, N. R., Eberlin, E. S., Mueser, K. T., & Correll, C. U. (2023). Efficacy and acceptability of psychosocial interventions in schizophrenia: Systematic overview and quality appraisal of the meta-analytic evidence. Molecular Psychiatry, 28(1), 354-368. https://doi.org/10.1038/s41380-022-01727-z",
        "Torres Hernández, R. C., González Lorenzo, C., & Martín Estévez, A. (2021). Eficacia de las técnicas cognitivas en la esquizofrenia. Revisión sistemática. Universidad de La Laguna. https://riull.ull.es/xmlui/handle/915/23977",
        "Vilar, T., Nogueira, M. J., Valentim, O., & Seabra, P. (2020). A psicoeducação na adesão terapêutica em utentes com esquizofrenia: Uma scoping review. Revista Portuguesa de Enfermagem de Saúde Mental, (spe7), 112-120. https://doi.org/10.19131/rpesm.0254",
        "World Health Organization. (2022, January 10). Schizophrenia. https://www.who.int/news-room/fact-sheets/detail/schizophrenia",
    ]
    for ref in references:
        add_reference(doc, ref)

    doc.save("/workspace/artigo_esquizofrenia_hospital_da_providencia.docx")


def build_cover_sheet():
    doc = Document()
    set_document_defaults(doc)
    doc.core_properties.title = "Folha de rosto - documento suplementar"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Folha de rosto - Documento Suplementar")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    add_paragraph(
        doc,
        "Título do manuscrito: Esquizofrenia e psicoterapia integrada: revisão teórica recente e análise das práticas públicas do P.A.I. Papa Francisco, Hospital São Francisco na Providência de Deus (Tijuca, Rio de Janeiro)",
        "Título do manuscrito:",
    )
    add_paragraph(doc, "Autor principal: [preencher nome completo, ORCID, afiliação, e-mail e telefone]", "Autor principal:")
    add_paragraph(doc, "Coautores, na ordem definitiva de autoria: [preencher nomes, ORCID, afiliações e e-mails]", "Coautores, na ordem definitiva de autoria:")
    add_paragraph(doc, "Autor correspondente: [preencher nome e endereço completo para correspondência]", "Autor correspondente:")
    add_paragraph(doc, "Contribuições de autoria (CRediT ou formato exigido pela revista): [preencher]", "Contribuições de autoria")
    add_paragraph(doc, "Conflitos de interesse: [declarar inexistência ou especificar]", "Conflitos de interesse:")
    add_paragraph(doc, "Financiamento: [declarar inexistência ou especificar agência/processo]", "Financiamento:")
    add_paragraph(doc, "Agradecimentos: [preencher somente se houver; não incluir no arquivo principal para preservar avaliação cega]", "Agradecimentos:")
    add_paragraph(doc, "Declaração de originalidade: Os autores declaram que a contribuição é original, inédita e não foi submetida simultaneamente a outra revista. Caso essa declaração não seja verdadeira no momento da submissão, justificar em 'Comentários ao Editor'.", "Declaração de originalidade:")
    add_paragraph(doc, "Declaração ética: Os autores declaram que o manuscrito de revisão não utilizou dados identificáveis de participantes humanos, prontuários, entrevistas ou imagens; foram consultadas apenas fontes bibliográficas e páginas institucionais públicas. Pesquisas empíricas futuras deverão ser submetidas a Comitê de Ética em Pesquisa.", "Declaração ética:")
    add_paragraph(doc, "Observação sobre autoria: Após a submissão, alterações na autoria podem não ser permitidas; a sequência registrada nesta folha de rosto deverá ser a sequência adotada para publicação, conforme normas da revista.", "Observação sobre autoria:")
    doc.save("/workspace/folha_de_rosto_documento_suplementar_modelo.docx")


def build_checklist():
    doc = Document()
    set_document_defaults(doc)
    doc.core_properties.title = "Checklist e declarações - documento suplementar"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Checklist de submissão e declarações - Documento Suplementar")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    add_paragraph(
        doc,
        "Atenção: este documento é um modelo operacional. O autor principal deve baixar o checklist oficial da revista escolhida, preencher, assinar, digitalizar e enviar o arquivo oficial como documento suplementar. Como a revista não foi especificada, os itens abaixo devem ser conferidos com as seções 'Preparação dos Originais' e 'Diretrizes para Autores' do periódico antes da submissão.",
    )
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Item")
    set_cell_text(hdr[1], "Status")
    set_cell_text(hdr[2], "Observação")
    items = [
        ("Arquivo principal em .doc/.docx, Times New Roman 12 e até 4 MB", "[ ] Conferido", "Arquivo gerado em .docx; verificar no Word antes de submeter."),
        ("Identificação de autoria removida do arquivo principal e propriedades do Word", "[ ] Conferido", "Metadados de autoria foram deixados em branco no arquivo gerado; revisar no Word."),
        ("Tema, palavras-chave e resumos em português, inglês e espanhol", "[ ] Conferido", "Incluídos no manuscrito."),
        ("Objetivo geral e quatro objetivos específicos", "[ ] Conferido", "Dois objetivos de revisão e dois relativos ao P.A.I./hospital."),
        ("Introdução, desenvolvimento, conclusão e referências", "[ ] Conferido", "Incluídos no manuscrito."),
        ("Citações e referências em APA 7ª edição", "[ ] Conferido", "Revisar conforme exigências específicas do periódico."),
        ("Declaração de originalidade e não submissão simultânea", "[ ] Assinar", "Se não for verdadeiro, justificar em Comentários ao Editor."),
        ("Declaração de cumprimento dos procedimentos éticos", "[ ] Assinar", "Revisão narrativa com fontes públicas; sem dados de participantes."),
        ("Folha de rosto preenchida como documento suplementar", "[ ] Preencher", "Usar modelo anexo ou formulário oficial da revista."),
        ("Checklist oficial da revista baixado, preenchido, assinado e digitalizado", "[ ] Pendente", "Depende da revista escolhida pelo autor principal."),
        ("Sequência de autores confirmada antes da submissão", "[ ] Confirmar", "Alterações posteriores podem não ser permitidas."),
    ]
    for item, status, obs in items:
        row = table.add_row().cells
        set_cell_text(row[0], item)
        set_cell_text(row[1], status)
        set_cell_text(row[2], obs)

    add_heading(doc, "Assinaturas", 1)
    add_paragraph(doc, "Autor principal: __________________________________________ Data: ____/____/______")
    add_paragraph(doc, "Coautores: _______________________________________________ Data: ____/____/______")
    doc.save("/workspace/checklist_declaracoes_documento_suplementar_modelo.docx")


if __name__ == "__main__":
    build_manuscript()
    build_cover_sheet()
    build_checklist()
