#!/usr/bin/env python3
"""Exporta BibTeX, RIS e DOCX (APA 7 básico)."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _bib_key(ref: dict[str, Any]) -> str:
    return re.sub(r"[^A-Za-z0-9]", "", ref.get("cite_key") or "ref")


def to_bibtex(refs: list[dict[str, Any]]) -> str:
    blocks = []
    for ref in refs:
        if not ref.get("located"):
            continue
        authors = " and ".join(
            a.replace(", ", ", ") for a in (ref.get("authors") or [])
        )
        fields = {
            "author": authors,
            "title": "{" + (ref.get("title") or "") + "}",
            "journal": ref.get("journal") or "",
            "year": str(ref.get("year") or ""),
            "volume": str(ref.get("volume") or ""),
            "number": str(ref.get("issue") or ""),
            "pages": str(ref.get("page") or ""),
            "doi": ref.get("doi") or "",
            "url": ref.get("url") or "",
        }
        body = ",\n".join(f"  {k} = {{{v}}}" for k, v in fields.items() if v and v != "None")
        blocks.append(f"@article{{{_bib_key(ref)},\n{body}\n}}")
    return "\n\n".join(blocks) + "\n"


def to_ris(refs: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    for ref in refs:
        if not ref.get("located"):
            continue
        lines.append("TY  - JOUR")
        for author in ref.get("authors") or []:
            # RIS prefers Last, First
            lines.append(f"AU  - {author}")
        lines.append(f"TI  - {ref.get('title') or ''}")
        if ref.get("journal"):
            lines.append(f"JO  - {ref['journal']}")
        if ref.get("year"):
            lines.append(f"PY  - {ref['year']}")
        if ref.get("volume"):
            lines.append(f"VL  - {ref['volume']}")
        if ref.get("issue") and str(ref["issue"]) not in {"0", "None"}:
            lines.append(f"IS  - {ref['issue']}")
        if ref.get("page"):
            lines.append(f"SP  - {ref['page']}")
        if ref.get("doi"):
            lines.append(f"DO  - {ref['doi']}")
        if ref.get("url"):
            lines.append(f"UR  - {ref['url']}")
        lines.append("ER  - ")
        lines.append("")
    return "\n".join(lines)


def _set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def _add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    _set_run_font(run, 12)


def markdown_to_docx(md_text: str, out_path: Path, title: str) -> None:
    """Converte markdown simples (見出し #/##, parágrafos) para DOCX APA-like."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    _add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)

    doc.core_properties.title = title
    doc.core_properties.author = "Antonio Augusto Dornelas de Andrade"
    doc.core_properties.language = "pt-BR"

    in_references = False
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run(line[2:].strip())
            _set_run_font(run, 12, bold=True)
            in_references = "referência" in line.lower()
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run(line[3:].strip())
            _set_run_font(run, 12, bold=True)
            in_references = "referência" in line.lower()
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run(line[4:].strip())
            _set_run_font(run, 12, bold=True)
            continue
        if line.startswith("|") and "---" not in line:
            # skip markdown tables in DOCX simple export; keep as paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line.replace("|", " ").strip())
            _set_run_font(run, 11)
            continue
        if line.strip().startswith("---"):
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if in_references:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.27)

        # bold **segments**
        parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                _set_run_font(run, 12, bold=True)
            else:
                # italics *segments*
                subparts = re.split(r"(\*[^*]+\*)", part)
                for sp in subparts:
                    if sp.startswith("*") and sp.endswith("*") and not sp.startswith("**"):
                        run = p.add_run(sp[1:-1])
                        _set_run_font(run, 12)
                        run.italic = True
                    else:
                        run = p.add_run(sp)
                        _set_run_font(run, 12)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
